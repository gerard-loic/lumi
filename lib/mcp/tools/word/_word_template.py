import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

from lib.config.config import Config

_ROOT_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

TABLE_STYLE_ID = "GabaritTableau"
TABLE_STYLE_NAME = "Tableau Gabarit"

SUMMARY_START_BOOKMARK = "LumiSommaireDebut"
SUMMARY_END_BOOKMARK = "LumiSommaireFin"
DOCUMENT_TITLE_BOOKMARK = "LumiDocTitre"

_COLOR_PRIMARY = "1F4E79"
_COLOR_ACCENT = "2E75B6"
_COLOR_BAND = "D9E6F2"
_COLOR_MUTED = "595959"


def template_path() -> str:
    relative = Config.get("word.template", default="templates/gabarit-document.docx")
    return relative if os.path.isabs(relative) else os.path.join(_ROOT_DIR, relative)


def placeholder_content() -> str:
    return Config.get("word.template_placeholder", default="{{content}}")


def placeholder_summary() -> str:
    return Config.get("word.template_summary_placeholder", default="{{summary}}")


def placeholder_title() -> str:
    return Config.get("word.template_title_placeholder", default="{{title}}")


def placeholder_date() -> str:
    return Config.get("word.template_date_placeholder", default="{{date}}")


def table_style_name() -> str:
    return Config.get("word.template_array_style", default=TABLE_STYLE_NAME)


def heading_style_name(level: int) -> str:
    return Config.get(f"word.heading_style_{level}", default=f"Heading {level}")


def page_break_before_heading1() -> bool:
    return Config.get("word.page_break_before_heading1", default=False)


def ensure_template() -> str:
    path = template_path()
    if not os.path.exists(path):
        build_template(path)
    return path


def _add_paragraph_style(styles, name, size, bold, color, italic=False):
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor.from_string(color)
    return style


def _configure_base_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in ((1, 18), (2, 15), (3, 13)):
        heading = styles[f"Heading {level}"]
        heading.font.name = "Calibri"
        heading.font.size = Pt(size)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor.from_string(_COLOR_PRIMARY)

    _add_paragraph_style(styles, "Titre Gabarit", size=32, bold=True, color=_COLOR_PRIMARY)
    _add_paragraph_style(styles, "Sous-titre Gabarit", size=18, bold=False, color=_COLOR_ACCENT)
    _add_paragraph_style(styles, "Date Gabarit", size=12, bold=False, color=_COLOR_MUTED, italic=True)
    _add_paragraph_style(styles, "Titre Sommaire", size=20, bold=True, color=_COLOR_PRIMARY)


def _configure_table_style(doc: Document) -> None:
    xml = f"""
    <w:style {nsdecls('w')} w:type="table" w:styleId="{TABLE_STYLE_ID}">
      <w:name w:val="{TABLE_STYLE_NAME}"/>
      <w:basedOn w:val="TableNormal"/>
      <w:uiPriority w:val="39"/>
      <w:tblPr>
        <w:tblStyleRowBandSize w:val="1"/>
        <w:tblBorders>
          <w:top w:val="single" w:sz="4" w:space="0" w:color="{_COLOR_PRIMARY}"/>
          <w:left w:val="single" w:sz="4" w:space="0" w:color="{_COLOR_PRIMARY}"/>
          <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{_COLOR_PRIMARY}"/>
          <w:right w:val="single" w:sz="4" w:space="0" w:color="{_COLOR_PRIMARY}"/>
          <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>
          <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>
        </w:tblBorders>
      </w:tblPr>
      <w:tblStylePr w:type="firstRow">
        <w:rPr>
          <w:b/>
          <w:color w:val="FFFFFF"/>
        </w:rPr>
        <w:tcPr>
          <w:shd w:val="clear" w:color="auto" w:fill="{_COLOR_PRIMARY}"/>
        </w:tcPr>
      </w:tblStylePr>
      <w:tblStylePr w:type="band1Horz">
        <w:tcPr>
          <w:shd w:val="clear" w:color="auto" w:fill="{_COLOR_BAND}"/>
        </w:tcPr>
      </w:tblStylePr>
    </w:style>
    """
    doc.styles.element.append(parse_xml(xml))


def _set_vertical_alignment(section, value="center") -> None:
    sectPr = section._sectPr
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), value)
    sectPr.insert_element_before(
        vAlign,
        "w:noEndnote", "w:titlePg", "w:textDirection", "w:bidi",
        "w:rtlGutter", "w:docGrid", "w:printerSettings", "w:sectPrChange",
    )


def _add_bottom_border(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), _COLOR_PRIMARY)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_simple_field(paragraph, instr: str, cached_text: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = cached_text
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def next_bookmark_id(doc: Document) -> int:
    ids = [int(e.get(qn("w:id"))) for e in doc.element.body.iter(qn("w:bookmarkStart"))]
    return max(ids, default=0) + 1


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    pPr.addnext(end)
    pPr.addnext(start)


def find_bookmark_paragraph_element(doc: Document, name: str):
    for start in doc.element.body.iter(qn("w:bookmarkStart")):
        if start.get(qn("w:name")) == name:
            return start.getparent()
    return None


def _enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.insert_element_before(
        update_fields,
        "w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr", "w:compat",
        "w:docVars", "w:rsids", "m:mathPr", "w:attachedSchema", "w:themeFontLang",
        "w:clrSchemeMapping", "w:doNotIncludeSubdocsInStats", "w:doNotAutoCompressPictures",
        "w:forceUpgrade", "w:captions", "w:readModeInkLockDown", "w:smartTagType",
        "sl:schemaLibrary", "w:shapeDefaults", "w:doNotEmbedSmartTags",
        "w:decimalSymbol", "w:listSeparator",
    )


def build_template(path: str | None = None) -> str:
    """
    Construit le gabarit Word par défaut (page de garde, en-tête/pied de page,
    zone de sommaire, styles de titres et de tableau) et l'enregistre sur disque.
    """
    path = path or template_path()

    doc = Document()
    _configure_base_styles(doc)
    _configure_table_style(doc)

    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin = section.bottom_margin = Cm(2)
    _set_vertical_alignment(section, "center")

    p_titre = doc.add_paragraph(style="Titre Gabarit")
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titre.add_run(placeholder_title())

    p_sous_titre = doc.add_paragraph(style="Sous-titre Gabarit")
    p_sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sous_titre.add_run(" ")

    for _ in range(8):
        doc.add_paragraph()

    p_date = doc.add_paragraph(style="Date Gabarit")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.add_run(placeholder_date())

    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    section2.left_margin = section2.right_margin = Cm(2.5)
    section2.top_margin = section2.bottom_margin = Cm(2)
    section2.header.is_linked_to_previous = False
    section2.footer.is_linked_to_previous = False

    header_p = section2.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run(placeholder_title())
    header_run.bold = True
    header_run.font.color.rgb = RGBColor.from_string(_COLOR_PRIMARY)
    _add_bottom_border(header_p)

    footer_p = section2.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("Page ")
    _add_simple_field(footer_p, "PAGE", "1")
    footer_p.add_run(" / ")
    _add_simple_field(footer_p, "NUMPAGES", "1")

    doc.add_paragraph("Sommaire", style="Titre Sommaire")
    marker_start = doc.add_paragraph()
    add_bookmark(marker_start, SUMMARY_START_BOOKMARK, 1)
    doc.add_paragraph(placeholder_summary())
    marker_end = doc.add_paragraph()
    add_bookmark(marker_end, SUMMARY_END_BOOKMARK, 2)

    doc.add_page_break()
    doc.add_paragraph(placeholder_content())

    _enable_update_fields(doc)

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_template())
