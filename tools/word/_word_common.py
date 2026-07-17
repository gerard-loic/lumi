import io
import re

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from tools.word._word_template import (
    DOCUMENT_TITLE_BOOKMARK,
    SUMMARY_END_BOOKMARK,
    SUMMARY_START_BOOKMARK,
    add_bookmark,
    find_bookmark_paragraph_element,
    heading_style_name,
    next_bookmark_id,
    page_break_before_heading1,
    placeholder_date,
    placeholder_summary,
    placeholder_title,
    table_style_name,
)

_BOOKMARK_PREFIX = "LumiTitre" #Utilisé pour identifier et générer les signets Word créés automatiquement pour chaque titre
_HTTPX_TIMEOUT = 15

_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"        # **bold**
    r"|(\*[^*]+\*)"            # *italic*
    r"|(~~[^~]+~~)"            # ~~strikethrough~~
    r"|(__[^_]+__)"            # __underline__
    r"|(`[^`]+`)"              # `inline code`
    r"|(\[[^\]]+\]\([^)]+\))"  # [text](url)
)


# ---------------------------------------------------------------------------
# Formatage inline / hyperliens
# ---------------------------------------------------------------------------

def write_inline(paragraph, text: str) -> None:
    last = 0
    for m in _INLINE_RE.finditer(text):
        start, end = m.start(), m.end()
        if start > last:
            paragraph.add_run(text[last:start])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("~~"):
            run = paragraph.add_run(token[2:-2])
            run.font.strike = True
        elif token.startswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.underline = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif token.startswith("["):
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_m:
                add_external_hyperlink(paragraph, link_m.group(1), link_m.group(2))
        last = end
    if last < len(text):
        paragraph.add_run(text[last:])


def add_external_hyperlink(paragraph, text: str, url: str) -> None:
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(paragraph, text: str, bookmark_name: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Détection des titres (Heading N)
# ---------------------------------------------------------------------------

def _style_labels(style) -> list[str]:
    """Noms sous lesquels un style peut être identifié : son nom, son id, et ses
    éventuels alias Word (ex. le panneau Styles de Word affiche parfois "Titre;diatem
    - titre 1", où "Titre" est le nom du style et "diatem - titre 1" un alias)."""
    labels = [style.name or "", style.style_id or ""]
    aliases_el = style.element.find(qn("w:aliases"))
    if aliases_el is not None:
        val = aliases_el.get(qn("w:val")) or ""
        labels.extend(a.strip() for a in re.split(r"[,;]", val) if a.strip())
    return [label for label in labels if label]


def _style_matches(style, configured: str | None) -> bool:
    if style is None or not configured:
        return False
    candidates = [p.strip() for p in configured.split(";") if p.strip()] or [configured]
    return not set(candidates).isdisjoint(_style_labels(style))


def heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style
    if style is None:
        return None
    for level in (1, 2, 3):
        if _style_matches(style, heading_style_name(level)) or _style_matches(style, f"Heading {level}"):
            return level
    return None


def find_heading_element(doc: Document, titre: str, case_sensitive: bool = False):
    needle = titre.strip() if case_sensitive else titre.strip().lower()
    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph) and heading_level(item) is not None:
            text = item.text.strip()
            cmp_text = text if case_sensitive else text.lower()
            if cmp_text == needle:
                return item
    return None


def find_section_end_element(doc: Document, heading_paragraph: Paragraph):
    """Retourne l'élément lxml du prochain titre de niveau <= celui de `heading_paragraph`,
    ou None si la section s'étend jusqu'à la fin du document."""
    level = heading_level(heading_paragraph)
    found = False
    for item in doc.iter_inner_content():
        if not found:
            if isinstance(item, Paragraph) and item._p is heading_paragraph._p:
                found = True
            continue
        if isinstance(item, Paragraph):
            lvl = heading_level(item)
            if lvl is not None and lvl <= level:
                return item._p
    return None


# ---------------------------------------------------------------------------
# Rendu Markdown -> Word
# ---------------------------------------------------------------------------

def _place(element, anchor) -> None:
    if anchor is not None:
        anchor.addprevious(element)


def _paragraph_with_style(doc: Document, style_name: str):
    """Crée un paragraphe avec `style_name` si ce style existe dans le document ;
    sinon un paragraphe simple. Retourne (paragraphe, style_applique: bool)."""
    try:
        return doc.add_paragraph(style=style_name), True
    except (KeyError, ValueError):
        return doc.add_paragraph(), False


def _find_style(doc: Document, configured: str | None):
    for style in doc.styles:
        if _style_matches(style, configured):
            return style
    return None


def _add_heading(doc: Document, text: str, level: int) -> Paragraph:
    """Ajoute un titre en utilisant le style configuré pour `level`
    (`word.heading_style_N`, reconnu par nom, id ou alias Word), avec repli sur
    le style Word natif "Heading N" si le style configuré n'existe pas dans le gabarit."""
    style = _find_style(doc, heading_style_name(level)) or _find_style(doc, f"Heading {level}")
    return doc.add_paragraph(text, style=style)


def _apply_table_style(table: Table, requested: str | None) -> None:
    for name in (requested, table_style_name(), "Table Grid"):
        if not name:
            continue
        try:
            table.style = name
            return
        except ValueError:
            continue


def render_table(doc: Document, rows: list[list[str]], table_style: str, anchor=None) -> Table:
    col_count = max((len(r) for r in rows), default=1)
    table = doc.add_table(rows=len(rows), cols=col_count)
    _apply_table_style(table, table_style)
    _place(table._tbl, anchor)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= col_count:
                continue
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    return table


def _parse_table(table_lines: list[str]) -> list[list[str]]:
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^:?-+:?$", c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


def render_markdown(
    doc: Document,
    markdown: str,
    chart_images: list,
    table_style: str,
    anchor=None,
    headings_out: list | None = None,
) -> None:
    """
    Rend le markdown enrichi (identique à la syntaxe historique de generer_fichier_word)
    dans `doc`. Si `anchor` (élément lxml) est fourni, chaque bloc généré est inséré juste
    avant cet élément ; sinon il est simplement ajouté à la fin du document.
    Chaque titre (# ## ###) rencontré reçoit un signet et est ajouté à `headings_out`
    sous la forme (niveau, texte, nom_du_signet) si cette liste est fournie.
    """
    bookmark_id = next_bookmark_id(doc)
    current_alignment = None
    pagebreak_before_h1 = page_break_before_heading1()
    any_content = False
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith(":::"):
            directive = line.strip()[3:].strip()
            if directive == "pagebreak":
                p = doc.add_page_break()
                _place(p._p, anchor)
                any_content = True
            elif directive == "center":
                current_alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif directive == "right":
                current_alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif directive == "justify":
                current_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif directive == "left":
                current_alignment = None
            elif directive.startswith("chart:"):
                idx = int(directive.split(":", 1)[1])
                if 0 <= idx < len(chart_images):
                    doc.add_picture(chart_images[idx], width=Inches(6))
                    _place(doc.paragraphs[-1]._p, anchor)
                    empty = doc.add_paragraph("")
                    _place(empty._p, anchor)
                    any_content = True
            elif directive.startswith("image:"):
                url = directive.split(":", 1)[1].strip()
                image = _fetch_image(url)
                if image is not None:
                    doc.add_picture(image, width=Inches(6))
                    _place(doc.paragraphs[-1]._p, anchor)
                    empty = doc.add_paragraph("")
                    _place(empty._p, anchor)
                    any_content = True
            i += 1
            continue

        if re.match(r"^\s*\|", line):
            table_lines = []
            while i < len(lines) and re.match(r"^\s*\|", lines[i]):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_table(table_lines)
            if rows:
                render_table(doc, rows, table_style, anchor=anchor)
                empty = doc.add_paragraph("")
                _place(empty._p, anchor)
                any_content = True
            continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            if current_alignment is not None:
                p.alignment = current_alignment
            _place(p._p, anchor)
            any_content = True
            i += 1
            continue

        p = None
        level = None
        if line.startswith("### "):
            p, level = _add_heading(doc, line[4:].strip(), 3), 3
        elif line.startswith("## "):
            p, level = _add_heading(doc, line[3:].strip(), 2), 2
        elif line.startswith("# "):
            if pagebreak_before_h1 and any_content:
                pb = doc.add_page_break()
                _place(pb._p, anchor)
            p, level = _add_heading(doc, line[2:].strip(), 1), 1
        elif re.match(r"^[-*] ", line):
            p, ok = _paragraph_with_style(doc, "List Bullet")
            write_inline(p, (line[2:].strip() if ok else f"• {line[2:].strip()}"))
        elif m := re.match(r"^(\d+)\. (.*)", line):
            p, ok = _paragraph_with_style(doc, "List Number")
            write_inline(p, (m.group(2).strip() if ok else f"{m.group(1)}. {m.group(2).strip()}"))
        elif line.strip() == "":
            p = doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            write_inline(p, line)

        if p is not None and current_alignment is not None:
            p.alignment = current_alignment

        if p is not None:
            _place(p._p, anchor)

        if p is not None and line.strip() != "":
            any_content = True

        if level is not None and headings_out is not None:
            bookmark_name = f"{_BOOKMARK_PREFIX}{bookmark_id}"
            add_bookmark(p, bookmark_name, bookmark_id)
            bookmark_id += 1
            headings_out.append((level, p.text.strip(), bookmark_name))

        i += 1


def _fetch_image(url: str) -> io.BytesIO | None:
    try:
        response = httpx.get(url, timeout=_HTTPX_TIMEOUT, follow_redirects=True)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Sommaire
# ---------------------------------------------------------------------------

def build_summary(doc: Document, headings: list[tuple[int, str, str]], anchor=None) -> None:
    if not headings:
        p = doc.add_paragraph("(Aucune section)")
        _place(p._p, anchor)
        return
    base_level = min(level for level, _, _ in headings)
    for level, text, bookmark_name in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18 * (level - base_level))
        _place(p._p, anchor)
        add_internal_hyperlink(p, text, bookmark_name)


def rebuild_summary(doc: Document, headings: list[tuple[int, str, str]]) -> None:
    start_p = find_bookmark_paragraph_element(doc, SUMMARY_START_BOOKMARK)
    end_p = find_bookmark_paragraph_element(doc, SUMMARY_END_BOOKMARK)

    if start_p is None or end_p is None:
        start_p, end_p = _mark_summary_zone(doc)
        if start_p is None or end_p is None:
            return

    node = start_p.getnext()
    while node is not None and node is not end_p:
        following = node.getnext()
        node.getparent().remove(node)
        node = following
    build_summary(doc, headings, anchor=end_p)


def _mark_summary_zone(doc: Document):
    """Si le gabarit ne définit pas encore de signets délimitant le sommaire, les crée
    autour du paragraphe {{summary}} (première génération) afin que les prochaines
    modifications puissent retrouver et reconstruire cette zone."""
    placeholder = find_placeholder_paragraph(doc, placeholder_summary())
    if placeholder is None:
        return None, None

    marker_start = placeholder.insert_paragraph_before()
    add_bookmark(marker_start, SUMMARY_START_BOOKMARK, next_bookmark_id(doc))

    marker_end = doc.add_paragraph()
    placeholder._p.addnext(marker_end._p)
    add_bookmark(marker_end, SUMMARY_END_BOOKMARK, next_bookmark_id(doc))

    placeholder._p.getparent().remove(placeholder._p)
    return marker_start._p, marker_end._p


def collect_headings(doc: Document) -> list[tuple[int, str, str]]:
    """Reconstruit la liste (niveau, texte, signet) des titres déjà présents dans le
    document, en réutilisant leur signet existant ou en leur en attribuant un nouveau."""
    headings = []
    bookmark_id = next_bookmark_id(doc)
    for item in doc.iter_inner_content():
        if not isinstance(item, Paragraph):
            continue
        level = heading_level(item)
        if level is None:
            continue
        existing = _paragraph_bookmark_name(item, prefix=_BOOKMARK_PREFIX)
        if existing is None:
            existing = f"{_BOOKMARK_PREFIX}{bookmark_id}"
            add_bookmark(item, existing, bookmark_id)
            bookmark_id += 1
        headings.append((level, item.text.strip(), existing))
    return headings


def _paragraph_bookmark_name(paragraph: Paragraph, prefix: str) -> str | None:
    for start in paragraph._p.findall(qn("w:bookmarkStart")):
        name = start.get(qn("w:name"))
        if name and name.startswith(prefix):
            return name
    return None


# ---------------------------------------------------------------------------
# Substitution de texte (placeholders et remplacements généraux)
# ---------------------------------------------------------------------------

def iter_all_paragraphs(doc: Document):
    """Parcourt tous les paragraphes du document, y compris ceux imbriqués dans des
    tableaux, zones de texte (text box) ou contrôles de contenu (SDT), dans le corps du
    document comme dans les en-têtes et pieds de page."""
    roots = [doc.element.body]
    for section in doc.sections:
        roots.append(section.header._element)
        roots.append(section.footer._element)

    seen = set()
    for root in roots:
        for p_el in root.iter(qn("w:p")):
            if p_el in seen:
                continue
            seen.add(p_el)
            yield Paragraph(p_el, doc)


_TEXT_RUN_CHILD_TAGS = {
    qn("w:t"),
    qn("w:tab"),
    qn("w:br"),
    qn("w:cr"),
    qn("w:noBreakHyphen"),
    qn("w:ptab"),
    qn("w:rPr"),
}


def _run_is_plain_text(run) -> bool:
    """Un run est "sûr" à réécrire s'il ne contient que du texte brut.

    Les runs de champs complexes (numérotation de page, etc.) sont composés de
    <w:fldChar>/<w:instrText> répartis sur plusieurs <w:r> frères dans le même
    paragraphe. python-docx expose ces runs avec un `.text` vide, donc les
    inclure dans la fusion/effacement ci-dessous détruirait le champ (ex. la
    pagination d'un en-tête) alors qu'ils ne contiennent aucun texte de
    substitution.
    """
    return all(child.tag in _TEXT_RUN_CHILD_TAGS for child in run._r)


def _replace_in_paragraph(paragraph: Paragraph, old: str, new: str, case_sensitive: bool) -> int:
    runs = [r for r in paragraph.runs if _run_is_plain_text(r)]
    if not runs or not old:
        return 0
    texts = [r.text for r in runs]
    full = "".join(texts)
    haystack = full if case_sensitive else full.lower()
    needle = old if case_sensitive else old.lower()

    spans = []
    search_pos = 0
    while True:
        found = haystack.find(needle, search_pos)
        if found == -1:
            break
        spans.append((found, found + len(old)))
        search_pos = found + len(needle)

    if not spans:
        return 0

    parts, last = [], 0
    for s, e in spans:
        parts.append(full[last:s])
        parts.append(new)
        last = e
    parts.append(full[last:])

    runs[0].text = "".join(parts)
    for r in runs[1:]:
        r.text = ""
    return len(spans)


def replace_text_everywhere(doc: Document, old: str, new: str, case_sensitive: bool = True) -> int:
    count = 0
    for paragraph in iter_all_paragraphs(doc):
        count += _replace_in_paragraph(paragraph, old, new, case_sensitive)
    return count


def replace_placeholders(doc: Document, titre: str, sous_titre: str, date_str: str) -> None:
    _bookmark_first_occurrence(doc, placeholder_title(), DOCUMENT_TITLE_BOOKMARK)
    replace_text_everywhere(doc, placeholder_title(), titre, case_sensitive=True)
    replace_text_everywhere(doc, placeholder_date(), date_str, case_sensitive=True)

    if not sous_titre:
        return
    for paragraph in doc.paragraphs:
        if paragraph.style is not None and paragraph.style.name == "Sous-titre Gabarit":
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = sous_titre
            else:
                paragraph.add_run(sous_titre)


def _bookmark_first_occurrence(doc: Document, text: str, bookmark_name: str) -> None:
    if find_bookmark_paragraph_element(doc, bookmark_name) is not None:
        return
    for p in iter_all_paragraphs(doc):
        if p.text.strip() == text:
            add_bookmark(p, bookmark_name, next_bookmark_id(doc))
            return


# ---------------------------------------------------------------------------
# Extraction (lecture) d'un document existant
# ---------------------------------------------------------------------------

def extract_outline(doc: Document) -> list[dict]:
    outline = []
    for p in doc.paragraphs:
        level = heading_level(p)
        if level is not None and p.text.strip():
            outline.append({"niveau": level, "titre": p.text.strip()})
    return outline


def extract_text(doc: Document) -> str:
    lines = []
    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            level = heading_level(item)
            lines.append(f"{'#' * level} {text}" if level else text)
        elif isinstance(item, Table):
            for row in item.rows:
                lines.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")
            lines.append("")
    return "\n".join(lines)


def extract_tables(doc: Document) -> list[dict]:
    results = []
    current_heading = None
    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            level = heading_level(item)
            if level is not None:
                current_heading = item.text.strip()
        elif isinstance(item, Table):
            results.append({
                "index": len(results),
                "section": current_heading,
                "lignes": [[cell.text for cell in row.cells] for row in item.rows],
            })
    return results


def find_placeholder_paragraph(doc: Document, text: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def extract_title(doc: Document) -> str:
    # Le signet peut avoir été préfixé par scope_bookmarks() lors d'une fusion de documents.
    for start in doc.element.body.iter(qn("w:bookmarkStart")):
        name = start.get(qn("w:name")) or ""
        if name == DOCUMENT_TITLE_BOOKMARK or name.endswith(f"_{DOCUMENT_TITLE_BOOKMARK}"):
            text = Paragraph(start.getparent(), doc).text.strip()
            if text:
                return text
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Titre Gabarit" and p.text.strip():
            return p.text.strip()
    return doc.core_properties.title or ""


def scope_bookmarks(doc: Document, scope: str) -> None:
    """Préfixe le nom de tous les signets du document (et les ancres internes qui les
    référencent) par `scope`, pour éviter toute collision lors d'une fusion de documents."""
    renames = {}
    for start in doc.element.body.iter(qn("w:bookmarkStart")):
        old_name = start.get(qn("w:name"))
        new_name = f"{scope}_{old_name}"
        renames[old_name] = new_name
        start.set(qn("w:name"), new_name)
    for hyperlink in doc.element.body.iter(qn("w:hyperlink")):
        anchor = hyperlink.get(qn("w:anchor"))
        if anchor in renames:
            hyperlink.set(qn("w:anchor"), renames[anchor])


def prepend_page_break(doc: Document) -> None:
    if not doc.paragraphs:
        return
    first = doc.paragraphs[0]
    break_p = first.insert_paragraph_before()
    break_p.add_run().add_break(WD_BREAK.PAGE)


def replace_table_data(doc: Document, index: int, rows: list[list[str]], table_style: str | None) -> Table:
    tables = doc.tables
    if index < 0 or index >= len(tables):
        raise ValueError(f"Aucun tableau à l'index {index} (le document en contient {len(tables)}).")
    old_table = tables[index]
    anchor = old_table._tbl.getnext()
    old_table._tbl.getparent().remove(old_table._tbl)
    return render_table(doc, rows, table_style, anchor=anchor)
