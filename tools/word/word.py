import io
import os
import shutil
import subprocess
import tempfile
from datetime import date
from typing import Annotated, Optional

from docx import Document
from pydantic import BaseModel, Field

from tools.charts._charts import render_chart
from tools.word import _word_common as wc
from tools.word._word_template import ensure_template, placeholder_content
from lib.agent.events import FileEvent
from lib.files.filestore import FileStore
from lib.mcp.tools import MCPTool, confirmation_tool

_MARKDOWN_DESC = (
    "Contenu du document en Markdown enrichi. "
    "Supporte : titres (# ## ###), gras (**texte**), italique (*texte*), "
    "barré (~~texte~~), souligné (__texte__), code inline (`texte`), "
    "liens ([texte](url)), tableaux (| col | col |) — limiter à 6 colonnes "
    "maximum pour un rendu correct dans la page ; au-delà, transposer le "
    "tableau (lignes ↔ colonnes) ou le scinder en plusieurs tableaux, "
    "listes à puces (- item) et numérotées (1. item), "
    "blocs de code (``` sur une ligne seule, puis code, puis ```). "
    "Directives de mise en page (sur leur propre ligne) : "
    ":::center, :::right, :::justify, :::left (alignement des blocs suivants), "
    ":::pagebreak (saut de page), "
    ":::chart:N (insère le graphique d'index N depuis le paramètre graphiques), "
    ":::image:URL (insère une image depuis une URL). "
    "Chaque titre (# ## ###) est automatiquement recensé dans le sommaire du document, "
    "avec un lien de navigation interne vers la section."
)

_GRAPHIQUES_DESC = (
    "Graphiques à insérer (liste d'objets). Chaque élément a : "
    "\"type\" ('barres', 'courbes' ou 'camembert'), "
    "\"titre\" (chaîne), "
    "\"données\" : pour 'barres'/'courbes' : {\"labels\":[...], \"series\":[{\"nom\":\"...\",\"valeurs\":[...]}]} ; "
    "pour 'camembert' : {\"labels\":[...], \"valeurs\":[...]}. "
    "Placer dans le markdown via :::chart:0, :::chart:1, etc. "
    "Exemple : [{\"type\":\"barres\",\"titre\":\"Ventes\","
    "\"données\":{\"labels\":[\"Jan\",\"Fév\"],\"series\":[{\"nom\":\"Produit A\",\"valeurs\":[10,20]}]}}]"
)


class FichierWord(BaseModel):
    url: str = Field(description="URL de téléchargement du fichier Word généré.")


class DocumentWordModifie(BaseModel):
    url: str = Field(description="URL de téléchargement du fichier Word mis à jour.")
    occurrences_remplacees: int = Field(description="Nombre d'occurrences remplacées dans le document.")


class SectionSommaire(BaseModel):
    niveau: int = Field(description="Niveau du titre (1, 2 ou 3).")
    titre: str = Field(description="Texte du titre.")


class ContenuDocumentWord(BaseModel):
    titre: str = Field(description="Titre du document (page de garde).")
    plan: list[SectionSommaire] = Field(description="Plan du document (liste des titres, dans l'ordre).")
    texte: str = Field(description="Texte intégral du document, titres préfixés par des dièses (#) selon leur niveau.")
    nombre_tableaux: int = Field(description="Nombre de tableaux présents dans le document.")


class TableauWord(BaseModel):
    index: int = Field(description="Index du tableau dans le document (0-based), à utiliser avec mettre_a_jour_tableau_word.")
    section: Optional[str] = Field(default=None, description="Titre de la section précédant ce tableau, si trouvée.")
    lignes: list[list[str]] = Field(description="Contenu du tableau, ligne par ligne (la première ligne est l'en-tête).")


class FichierPDF(BaseModel):
    url: str = Field(description="URL de téléchargement du fichier PDF généré.")


class WordService(MCPTool):
    name = "word"
    description = "Génération et modification de fichiers Word respectant le gabarit de l'entreprise"

    @confirmation_tool(
        question="Je vais générer le fichier Word. Dois-je continuer ?",
        options=["Oui", "Non"],
        validation_option=0,
    )
    def generer_fichier_word(
        self,
        titre: Annotated[str, Field(description="Titre du document, affiché sur la page de garde et dans l'en-tête de chaque page.")],
        contenu_markdown: Annotated[str, Field(description=_MARKDOWN_DESC)],
        sous_titre: Annotated[
            Optional[str],
            Field(default=None, description="Sous-titre affiché sous le titre sur la page de garde."),
        ] = None,
        graphiques: Annotated[Optional[list[dict]], Field(default=None, description=_GRAPHIQUES_DESC)] = None,
        style_tableau: Annotated[
            Optional[str],
            Field(default=None, description="Nom d'un style de tableau du gabarit à utiliser. Si omis, le style du gabarit est utilisé."),
        ] = None,
        nom_fichier: Annotated[
            Optional[str],
            Field(default=None, description="Nom du fichier Word (ex: 'rapport.docx'). Si omis, 'document.docx' est utilisé."),
        ] = None,
    ) -> FichierWord:
        """
        Génère un fichier Word (.docx) respectant le gabarit de l'entreprise : page de garde
        (titre, sous-titre, date), en-tête et pied de page avec le titre et la pagination,
        sommaire cliquable généré automatiquement à partir des titres du contenu, et styles
        de tableau du gabarit. À utiliser dès que l'utilisateur demande un export Word, un
        rapport, un compte-rendu ou un document téléchargeable au format Word.
        Supporte l'insertion de graphiques (barres, courbes, camembert) via le paramètre graphiques.
        """
        filename = (nom_fichier or "document").removesuffix(".docx") + ".docx"
        doc = Document(ensure_template())

        marker = wc.find_placeholder_paragraph(doc, placeholder_content())
        if marker is not None:
            marker._p.getparent().remove(marker._p)

        chart_images = [render_chart(c) for c in (graphiques or [])]

        headings: list = []
        wc.render_markdown(
            doc, contenu_markdown, chart_images,
            style_tableau, anchor=None, headings_out=headings,
        )
        wc.rebuild_summary(doc, headings)
        wc.replace_placeholders(doc, titre, sous_titre or "", date.today().strftime("%d/%m/%Y"))

        return self._save(doc, filename)

    def lire_document_word(
        self,
        url_document: Annotated[str, Field(description="URL du fichier Word à lire (URL de téléchargement retournée par un outil Word précédent).")],
    ) -> ContenuDocumentWord:
        """
        Lit un fichier Word précédemment généré et retourne son titre, son plan (liste des
        titres) et son texte intégral (les titres étant préfixés par des dièses selon leur
        niveau). À utiliser avant de modifier un document pour connaître son contenu et
        savoir où insérer ou remplacer des éléments (ex : les titres exacts pour ajouter_contenu_word).
        """
        doc = self._load(url_document)
        return ContenuDocumentWord(
            titre=wc.extract_title(doc),
            plan=[SectionSommaire(**o) for o in wc.extract_outline(doc)],
            texte=wc.extract_text(doc),
            nombre_tableaux=len(doc.tables),
        )

    def extraire_tableaux_word(
        self,
        url_document: Annotated[str, Field(description="URL du fichier Word contenant les tableaux à extraire.")],
    ) -> list[TableauWord]:
        """
        Extrait tous les tableaux d'un document Word, avec leur index (à utiliser avec
        mettre_a_jour_tableau_word) et la section qui les précède. À utiliser pour connaître
        le contenu exact d'un tableau avant de le modifier.
        """
        doc = self._load(url_document)
        return [TableauWord(**t) for t in wc.extract_tables(doc)]

    @confirmation_tool(
        question="Je vais remplacer ce texte dans le document. Dois-je continuer ?",
        options=["Oui", "Non"],
        validation_option=0,
    )
    def remplacer_texte_word(
        self,
        url_document: Annotated[str, Field(description="URL du fichier Word à modifier.")],
        texte_recherche: Annotated[str, Field(description="Texte exact à rechercher.")],
        texte_remplacement: Annotated[str, Field(description="Texte de remplacement.")],
        respecter_casse: Annotated[bool, Field(default=True, description="Si faux, la recherche ignore la casse.")] = True,
        nom_fichier: Annotated[
            Optional[str],
            Field(default=None, description="Nom du fichier Word résultant. Si omis, le nom du document d'origine est repris."),
        ] = None,
    ) -> DocumentWordModifie:
        """
        Recherche et remplace toutes les occurrences d'un texte dans un document Word
        existant (corps du texte, en-têtes, pieds de page et tableaux compris) et retourne
        le fichier modifié. À utiliser pour corriger ou mettre à jour un texte précis
        (y compris le titre affiché en en-tête, si celui-ci a été inséré via generer_fichier_word).
        """
        doc = self._load(url_document)
        count = wc.replace_text_everywhere(doc, texte_recherche, texte_remplacement, case_sensitive=respecter_casse)
        filename = nom_fichier or self._url_filename(url_document, "document-modifie.docx")
        result = self._save(doc, filename)
        return DocumentWordModifie(url=result.url, occurrences_remplacees=count)

    @confirmation_tool(
        question="Je vais ajouter ce contenu au document. Dois-je continuer ?",
        options=["Oui", "Non"],
        validation_option=0,
    )
    def ajouter_contenu_word(
        self,
        url_document: Annotated[str, Field(description="URL du fichier Word à compléter.")],
        contenu_markdown: Annotated[str, Field(description=_MARKDOWN_DESC)],
        graphiques: Annotated[Optional[list[dict]], Field(default=None, description=_GRAPHIQUES_DESC)] = None,
        position: Annotated[
            str,
            Field(
                default="fin",
                description=(
                    "Où insérer le contenu : 'fin' pour l'ajouter à la fin du document (par défaut), "
                    "ou 'apres:<titre exact>' pour l'insérer à la fin de la section portant ce titre "
                    "(avant le prochain titre de même niveau ou de niveau supérieur). "
                    "Utiliser lire_document_word pour connaître les titres exacts disponibles."
                ),
            ),
        ] = "fin",
        nom_fichier: Annotated[
            Optional[str],
            Field(default=None, description="Nom du fichier Word résultant. Si omis, le nom du document d'origine est repris."),
        ] = None,
    ) -> FichierWord:
        """
        Ajoute du contenu Markdown enrichi (texte, titres, listes, tableaux, graphiques) à un
        document Word existant, à la fin du document ou dans une section précise, et
        recalcule automatiquement le sommaire. À utiliser pour compléter un document déjà
        généré sans le reconstruire entièrement.
        """
        doc = self._load(url_document)
        chart_images = [render_chart(c) for c in (graphiques or [])]

        anchor = None
        if position and position.strip().lower().startswith("apres:"):
            titre_section = position.split(":", 1)[1].strip()
            heading = wc.find_heading_element(doc, titre_section)
            if heading is None:
                raise ValueError(f"Section introuvable dans le document : « {titre_section} ».")
            anchor = wc.find_section_end_element(doc, heading)

        wc.render_markdown(doc, contenu_markdown, chart_images, None, anchor=anchor, headings_out=[])
        wc.rebuild_summary(doc, wc.collect_headings(doc))

        filename = nom_fichier or self._url_filename(url_document, "document-modifie.docx")
        return self._save(doc, filename)

    @confirmation_tool(
        question="Je vais mettre à jour ce tableau. Dois-je continuer ?",
        options=["Oui", "Non"],
        validation_option=0,
    )
    def mettre_a_jour_tableau_word(
        self,
        url_document: Annotated[str, Field(description="URL du fichier Word contenant le tableau à mettre à jour.")],
        index_tableau: Annotated[int, Field(description="Index du tableau à remplacer (0-based), obtenu via extraire_tableaux_word.")],
        lignes: Annotated[
            list[list[str]],
            Field(description="Nouvelles lignes du tableau (la première ligne doit être l'en-tête). Remplace entièrement le tableau existant. Limiter à 6 colonnes maximum pour un rendu correct dans la page ; au-delà, transposer le tableau (lignes ↔ colonnes) ou le scinder en plusieurs tableaux."),
        ],
        style_tableau: Annotated[
            Optional[str],
            Field(default=None, description="Nom d'un style de tableau du gabarit à appliquer. Si omis, le style du gabarit est utilisé."),
        ] = None,
        nom_fichier: Annotated[
            Optional[str],
            Field(default=None, description="Nom du fichier Word résultant. Si omis, le nom du document d'origine est repris."),
        ] = None,
    ) -> FichierWord:
        """
        Remplace intégralement les données d'un tableau existant dans un document Word
        (identifié par son index, voir extraire_tableaux_word) tout en conservant sa position
        et le style de tableau du gabarit. À utiliser pour mettre à jour des chiffres ou des
        lignes dans un tableau déjà présent dans un document généré.
        """
        doc = self._load(url_document)
        wc.replace_table_data(doc, index_tableau, lignes, style_tableau)
        filename = nom_fichier or self._url_filename(url_document, "document-modifie.docx")
        return self._save(doc, filename)

    @confirmation_tool(
        question="Je vais fusionner ces documents Word. Dois-je continuer ?",
        options=["Oui", "Non"],
        validation_option=0,
    )
    def fusionner_documents_word(
        self,
        urls_documents: Annotated[
            list[str],
            Field(description="Liste ordonnée d'au moins deux URL de fichiers Word à fusionner en un seul document."),
        ],
        nom_fichier: Annotated[
            Optional[str],
            Field(default=None, description="Nom du fichier Word fusionné (ex: 'dossier-complet.docx'). Si omis, 'document-fusionne.docx' est utilisé."),
        ] = None,
    ) -> FichierWord:
        """
        Fusionne plusieurs documents Word en un seul fichier, dans l'ordre fourni, en
        insérant un saut de page entre chaque document. Chaque document conserve sa propre
        page de garde et son propre sommaire. À utiliser pour assembler plusieurs rapports
        ou sections générées séparément en un seul livrable.
        """
        if len(urls_documents) < 2:
            raise ValueError("Il faut fournir au moins deux documents à fusionner.")

        from docxcompose.composer import Composer

        docs = [self._load(url) for url in urls_documents]
        for i, doc in enumerate(docs):
            wc.scope_bookmarks(doc, f"doc{i}")
        for doc in docs[1:]:
            wc.prepend_page_break(doc)

        composer = Composer(docs[0])
        for doc in docs[1:]:
            composer.append(doc)

        filename = (nom_fichier or "document-fusionne").removesuffix(".docx") + ".docx"
        return self._save(composer.doc, filename)

    @confirmation_tool(
        question="Je vais convertir ce fichier Word en PDF. Dois-je continuer ?",
        options=["Oui", "Non"],
        validation_option=0,
    )
    def convertir_word_en_pdf(
        self,
        url_document: Annotated[str, Field(description="URL du fichier Word à convertir en PDF.")],
        nom_fichier: Annotated[
            Optional[str],
            Field(default=None, description="Nom du fichier PDF résultant. Si omis, le nom du document d'origine est repris."),
        ] = None,
    ) -> FichierPDF:
        """
        Convertit un document Word existant en PDF, en conservant fidèlement sa mise en page
        (police, tableaux, en-têtes/pieds de page). À utiliser quand l'utilisateur demande la
        version PDF d'un document Word déjà généré.
        """
        content = FileStore.load(url_document)
        pdf_bytes = _convert_docx_to_pdf(content)

        base_name = os.path.splitext(nom_fichier or FileStore.filename_from_url(url_document) or "document")[0]
        filename = f"{base_name}.pdf"

        url = FileStore.save(filename=filename, content=pdf_bytes)
        self.emit(FileEvent.get(name=filename, url=url))
        return FichierPDF(url=url)

    def _save(self, doc: Document, filename: str) -> FichierWord:
        buffer = io.BytesIO()
        doc.save(buffer)
        url = FileStore.save(filename=filename, content=buffer.getvalue())
        self.emit(FileEvent.get(name=filename, url=url))
        return FichierWord(url=url)

    def _load(self, url_document: str) -> Document:
        return Document(io.BytesIO(FileStore.load(url_document)))

    def _url_filename(self, url_document: str, fallback: str) -> str:
        return FileStore.filename_from_url(url_document) or fallback


def _find_soffice() -> str | None:
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path
    for candidate in ("/usr/bin/soffice", "/usr/bin/libreoffice", "/opt/libreoffice/program/soffice"):
        if os.path.exists(candidate):
            return candidate
    return None


def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError(
            "La conversion Word vers PDF nécessite LibreOffice (commande 'soffice' ou "
            "'libreoffice') installé sur le serveur, ce qui n'est pas le cas ici."
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        src_path = os.path.join(tmpdir, "document.docx")
        with open(src_path, "wb") as f:
            f.write(docx_bytes)

        subprocess.run(
            [soffice, "--headless", "--norestore", "--convert-to", "pdf", "--outdir", tmpdir, src_path],
            check=True, capture_output=True, timeout=120,
        )

        pdf_path = os.path.join(tmpdir, "document.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("La conversion en PDF a échoué : aucun fichier PDF produit.")
        with open(pdf_path, "rb") as f:
            return f.read()
