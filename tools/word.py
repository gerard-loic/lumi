import io
import re
from typing import Annotated, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pydantic import BaseModel, Field

from tools._charts import render_chart
from lib.agent.events import FileEvent
from lib.files.filestore import FileStore
from lib.mcp.tools import MCPTool, confirmation_tool, native_tool


class FichierWord(BaseModel):
    url: str = Field(description="URL de téléchargement du fichier Word généré.")


class WordService(MCPTool):
    name = "word"
    description = "Génération de fichiers Word"

    @native_tool
    @confirmation_tool(
        question="Je vais générer le fichier Word. Dois-je continuer ?",
        options=["Oui", "Non"],
        validation_option=0,
    )
    def generer_fichier_word(
        self,
        contenu_markdown: Annotated[
            str,
            Field(
                description=(
                    "Contenu du document en Markdown enrichi. "
                    "Supporte : titres (# ## ###), gras (**texte**), italique (*texte*), "
                    "barré (~~texte~~), souligné (__texte__), code inline (`texte`), "
                    "liens ([texte](url)), tableaux (| col | col |), "
                    "listes à puces (- item) et numérotées (1. item), "
                    "blocs de code (``` sur une ligne seule, puis code, puis ```). "
                    "Directives de mise en page (sur leur propre ligne) : "
                    ":::center, :::right, :::justify, :::left (alignement des blocs suivants), "
                    ":::pagebreak (saut de page), "
                    ":::chart:N (insère le graphique d'index N depuis le paramètre graphiques)."
                )
            ),
        ],
        graphiques: Annotated[
            Optional[list[dict]],
            Field(
                default=None,
                description=(
                    "Graphiques à insérer (liste d'objets). Chaque élément a : "
                    "\"type\" ('barres', 'courbes' ou 'camembert'), "
                    "\"titre\" (chaîne), "
                    "\"données\" : pour 'barres'/'courbes' : {\"labels\":[...], \"series\":[{\"nom\":\"...\",\"valeurs\":[...]}]} ; "
                    "pour 'camembert' : {\"labels\":[...], \"valeurs\":[...]}. "
                    "Placer dans le markdown via :::chart:0, :::chart:1, etc. "
                    "Exemple : [{\"type\":\"barres\",\"titre\":\"Ventes\","
                    "\"données\":{\"labels\":[\"Jan\",\"Fév\"],\"series\":[{\"nom\":\"Produit A\",\"valeurs\":[10,20]}]}}]"
                ),
            ),
        ] = None,
        nom_fichier: Annotated[
            Optional[str],
            Field(
                default=None,
                description="Nom du fichier Word (ex: 'rapport.docx'). Si omis, 'document.docx' est utilisé.",
            ),
        ] = None,
    ) -> FichierWord:
        """
        Génère un fichier Word (.docx) à partir de contenu Markdown enrichi et retourne l'URL de téléchargement.
        À utiliser dès que l'utilisateur demande un export Word, un document téléchargeable, un rapport ou un compte-rendu au format Word.
        Supporte l'insertion de graphiques (barres, courbes, camembert) via le paramètre graphiques.
        """
        filename = (nom_fichier or "document").removesuffix(".docx") + ".docx"
        doc = Document()

        charts = graphiques or []
        chart_images = [render_chart(c) for c in charts]

        current_alignment = None
        lines = contenu_markdown.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            if line.strip().startswith(":::"):
                directive = line.strip()[3:].strip()
                if directive == "pagebreak":
                    doc.add_page_break()
                elif directive == "center":
                    current_alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif directive == "right":
                    current_alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif directive == "justify":
                    current_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif directive == "left":
                    current_alignment = None
                elif directive.startswith("chart:"):
                    idx = int(directive.split(":", 1)[1])
                    if 0 <= idx < len(chart_images):
                        doc.add_picture(chart_images[idx], width=Inches(6))
                        doc.add_paragraph("")
                i += 1
                continue

            if re.match(r"^\s*\|", line):
                table_lines = []
                while i < len(lines) and re.match(r"^\s*\|", lines[i]):
                    table_lines.append(lines[i])
                    i += 1
                self._render_table(doc, _parse_table(table_lines))
                continue

            if line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                if current_alignment is not None:
                    p.alignment = current_alignment
                i += 1
                continue

            p = None
            if line.startswith("### "):
                p = doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                p = doc.add_heading(line[2:].strip(), level=1)
            elif re.match(r"^[-*] ", line):
                p = doc.add_paragraph(style="List Bullet")
                _write_inline(p, line[2:].strip())
            elif m := re.match(r"^(\d+)\. (.*)", line):
                p = doc.add_paragraph(style="List Number")
                _write_inline(p, m.group(2).strip())
            elif line.strip() == "":
                p = doc.add_paragraph("")
            else:
                p = doc.add_paragraph()
                _write_inline(p, line)

            if p is not None and current_alignment is not None:
                p.alignment = current_alignment

            i += 1

        buffer = io.BytesIO()
        doc.save(buffer)

        url = FileStore.save(filename=filename, content=buffer.getvalue())
        self.emit(FileEvent.get(name=filename, url=url))
        return FichierWord(url=url)

    def _render_table(self, doc: Document, rows: list[list[str]]) -> None:
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph("")


_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"        # **bold**
    r"|(\*[^*]+\*)"            # *italic*
    r"|(~~[^~]+~~)"            # ~~strikethrough~~
    r"|(__[^_]+__)"            # __underline__
    r"|(`[^`]+`)"              # `inline code`
    r"|(\[[^\]]+\]\([^)]+\))"  # [text](url)
)


def _write_inline(paragraph, text: str) -> None:
    last = 0
    for m in _INLINE_RE.finditer(text):
        start, end = m.start(), m.end()
        if start > last:
            paragraph.add_run(text[last:start])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("~~"):
            run = paragraph.add_run(token[2:-2])
            run.font.strike = True
        elif token.startswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.underline = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif token.startswith("["):
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_m:
                _add_hyperlink(paragraph, link_m.group(1), link_m.group(2))
        last = end
    if last < len(text):
        paragraph.add_run(text[last:])


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _parse_table(table_lines: list[str]) -> list[list[str]]:
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^:?-+:?$", c) for c in cells if c):
            continue
        rows.append(cells)
    return rows
